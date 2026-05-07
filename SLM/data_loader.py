"""
Enterprise Data Loader and Preprocessor
========================================
Handles loading and preprocessing of heterogeneous enterprise data from multiple sources.
Supports: documents, code, chat logs, transcripts, and structured data.
"""

import os
import re
import json
import logging
from pathlib import Path
from typing import List, Dict, Iterator, Optional, Tuple
from dataclasses import dataclass
import unicodedata

# Document processing libraries
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("Warning: PyPDF2 not installed. PDF support disabled.")

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. DOCX support disabled.")

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False
    print("Warning: pandas not installed. Excel/CSV support limited.")

from config import (
    RAW_DATA_DIR, 
    PROCESSED_DATA_DIR,
    PREPROCESSING_CONFIG,
    DATA_SOURCES,
)

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


# ============================================================================
# DATA CLASSES
# ============================================================================
@dataclass
class Document:
    """Container for a processed document."""
    text: str
    metadata: Dict
    doc_type: str  # 'document', 'code', 'chat', 'transcript'
    source: str
    
    def __len__(self):
        return len(self.text)
    
    def to_dict(self):
        return {
            'text': self.text,
            'metadata': self.metadata,
            'doc_type': self.doc_type,
            'source': self.source,
        }


# ============================================================================
# TEXT CLEANING AND NORMALIZATION
# ============================================================================
class TextCleaner:
    """Utility class for cleaning and normalizing text."""
    
    @staticmethod
    def remove_html_tags(text: str) -> str:
        """Remove HTML tags from text."""
        return re.sub(r'<[^>]+>', '', text)
    
    @staticmethod
    def remove_urls(text: str) -> str:
        """Remove URLs from text."""
        return re.sub(r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+', '', text)
    
    @staticmethod
    def remove_emails(text: str) -> str:
        """Remove email addresses from text."""
        return re.sub(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', '', text)
    
    @staticmethod
    def fix_unicode(text: str) -> str:
        """Fix unicode encoding issues."""
        return unicodedata.normalize('NFKC', text)
    
    @staticmethod
    def normalize_whitespace(text: str) -> str:
        """Normalize whitespace to single spaces."""
        text = re.sub(r'\s+', ' ', text)
        text = text.strip()
        return text
    
    @staticmethod
    def clean_text(text: str, config: dict = None) -> str:
        """
        Apply all cleaning operations based on configuration.
        """
        if config is None:
            config = PREPROCESSING_CONFIG
        
        if config.get('remove_html', True):
            text = TextCleaner.remove_html_tags(text)
        
        if config.get('remove_urls', False):
            text = TextCleaner.remove_urls(text)
        
        if config.get('remove_emails', False):
            text = TextCleaner.remove_emails(text)
        
        if config.get('fix_unicode', True):
            text = TextCleaner.fix_unicode(text)
        
        if config.get('normalize_whitespace', True):
            text = TextCleaner.normalize_whitespace(text)
        
        return text


# ============================================================================
# FILE READERS
# ============================================================================
class FileReader:
    """Read and extract text from various file formats."""
    
    @staticmethod
    def read_text(file_path: Path) -> str:
        """Read plain text file."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        except Exception as e:
            logger.error(f"Error reading text file {file_path}: {e}")
            return ""
    
    @staticmethod
    def read_pdf(file_path: Path) -> str:
        """Read PDF file."""
        if not PDF_AVAILABLE:
            logger.warning(f"Cannot read PDF {file_path}: PyPDF2 not installed")
            return ""
        
        try:
            text = []
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page in pdf_reader.pages:
                    text.append(page.extract_text())
            return '\n'.join(text)
        except Exception as e:
            logger.error(f"Error reading PDF {file_path}: {e}")
            return ""
    
    @staticmethod
    def read_docx(file_path: Path) -> str:
        """Read DOCX file."""
        if not DOCX_AVAILABLE:
            logger.warning(f"Cannot read DOCX {file_path}: python-docx not installed")
            return ""
        
        try:
            doc = DocxDocument(file_path)
            return '\n'.join([para.text for para in doc.paragraphs])
        except Exception as e:
            logger.error(f"Error reading DOCX {file_path}: {e}")
            return ""
    
    @staticmethod
    def read_json(file_path: Path) -> str:
        """Read JSON file and convert to text."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            return json.dumps(data, indent=2)
        except Exception as e:
            logger.error(f"Error reading JSON {file_path}: {e}")
            return ""
    
    @staticmethod
    def read_csv(file_path: Path) -> str:
        """Read CSV file and convert to text."""
        if not PANDAS_AVAILABLE:
            return FileReader.read_text(file_path)
        
        try:
            df = pd.read_csv(file_path)
            return df.to_string()
        except Exception as e:
            logger.error(f"Error reading CSV {file_path}: {e}")
            return ""
    
    @staticmethod
    def read_file(file_path: Path) -> Tuple[str, str]:
        """
        Read file and return text content with detected type.
        Returns: (text, doc_type)
        """
        suffix = file_path.suffix.lower()
        
        # Detect document type
        code_extensions = {'.py', '.js', '.java', '.cpp', '.c', '.h', '.html', '.css', '.go', '.rs'}
        if suffix in code_extensions:
            doc_type = 'code'
        else:
            doc_type = 'document'
        
        # Read based on extension
        if suffix == '.pdf':
            text = FileReader.read_pdf(file_path)
        elif suffix in {'.docx', '.doc'}:
            text = FileReader.read_docx(file_path)
        elif suffix == '.json':
            text = FileReader.read_json(file_path)
        elif suffix == '.csv':
            text = FileReader.read_csv(file_path)
        else:
            text = FileReader.read_text(file_path)
        
        return text, doc_type


# ============================================================================
# DOCUMENT PROCESSOR
# ============================================================================
class DocumentProcessor:
    """Process documents and split into chunks."""
    
    def __init__(self, config: dict = None):
        self.config = config or PREPROCESSING_CONFIG
        self.cleaner = TextCleaner()
    
    def chunk_document(self, text: str, max_length: int = None, overlap: int = None) -> List[str]:
        """
        Split document into overlapping chunks.
        """
        if max_length is None:
            max_length = self.config['max_document_length']
        if overlap is None:
            overlap = self.config['document_overlap']
        
        if len(text) <= max_length:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + max_length
            chunk = text[start:end]
            
            # Try to break at sentence boundary
            if end < len(text):
                last_period = chunk.rfind('.')
                last_newline = chunk.rfind('\n')
                break_point = max(last_period, last_newline)
                
                if break_point > 0:
                    chunk = chunk[:break_point + 1]
                    end = start + break_point + 1
            
            chunks.append(chunk.strip())
            start = end - overlap
            
            if start >= len(text):
                break
        
        return chunks
    
    def process_document(
        self,
        file_path: Path,
        source: str = "local"
    ) -> List[Document]:
        """
        Process a single document file.
        Returns list of Document objects (may be chunked).
        """
        # Check if file type is supported
        if file_path.suffix.lower() not in self.config['supported_extensions']:
            logger.warning(f"Unsupported file type: {file_path}")
            return []
        
        # Read file
        text, doc_type = FileReader.read_file(file_path)
        
        if not text or len(text) < self.config['min_document_length']:
            logger.warning(f"Document too short or empty: {file_path}")
            return []
        
        # Clean text
        text = self.cleaner.clean_text(text, self.config)
        
        # Chunk if necessary
        chunks = self.chunk_document(text)
        
        # Create Document objects
        documents = []
        for i, chunk in enumerate(chunks):
            doc = Document(
                text=chunk,
                metadata={
                    'file_path': str(file_path),
                    'file_name': file_path.name,
                    'chunk_id': i,
                    'total_chunks': len(chunks),
                    'file_size': file_path.stat().st_size,
                },
                doc_type=doc_type,
                source=source,
            )
            documents.append(doc)
        
        return documents


# ============================================================================
# DATA LOADER
# ============================================================================
class EnterpriseDataLoader:
    """
    Main data loader for enterprise sources.
    Aggregates documents from multiple sources and formats.
    """
    
    def __init__(self, config: dict = None):
        self.config = config or DATA_SOURCES
        self.processor = DocumentProcessor()
        self.documents = []
    
    def load_local_files(self, paths: List[str]) -> List[Document]:
        """Load documents from local file system."""
        documents = []
        
        for path_str in paths:
            path = Path(path_str)
            
            if not path.exists():
                logger.warning(f"Path does not exist: {path}")
                continue
            
            # Recursively find all files
            if path.is_dir():
                files = []
                for ext in PREPROCESSING_CONFIG['supported_extensions']:
                    files.extend(path.rglob(f"*{ext}"))
            else:
                files = [path]
            
            logger.info(f"Found {len(files)} files in {path}")
            
            # Process each file
            for file_path in files:
                try:
                    docs = self.processor.process_document(file_path, source="local")
                    documents.extend(docs)
                except Exception as e:
                    logger.error(f"Error processing {file_path}: {e}")
        
        return documents
    
    def load_sharepoint(self, config: dict) -> List[Document]:
        """
        Load documents from SharePoint.
        NOTE: Requires additional setup and authentication.
        """
        logger.warning("SharePoint integration requires Office365-REST-Python-Client library")
        logger.warning("Please implement authentication and file download logic")
        # Placeholder for SharePoint integration
        return []
    
    def load_google_drive(self, config: dict) -> List[Document]:
        """
        Load documents from Google Drive.
        NOTE: Requires Google Drive API setup and credentials.
        """
        logger.warning("Google Drive integration requires google-api-python-client library")
        logger.warning("Please implement authentication and file download logic")
        # Placeholder for Google Drive integration
        return []
    
    def load_slack(self, config: dict) -> List[Document]:
        """
        Load chat logs from Slack.
        NOTE: Requires Slack API token and slack-sdk library.
        """
        logger.warning("Slack integration requires slack-sdk library")
        logger.warning("Please implement Slack API logic")
        # Placeholder for Slack integration
        return []
    
    def load_github(self, config: dict) -> List[Document]:
        """
        Load code repositories from GitHub.
        NOTE: Requires PyGithub library.
        """
        logger.warning("GitHub integration requires PyGithub library")
        logger.warning("Please implement GitHub API logic")
        # Placeholder for GitHub integration
        return []
    
    def load_all_sources(self) -> List[Document]:
        """Load documents from all enabled sources."""
        all_documents = []
        
        # Local files
        if self.config['local_files']['enabled']:
            logger.info("Loading local files...")
            docs = self.load_local_files(self.config['local_files']['paths'])
            all_documents.extend(docs)
            logger.info(f"Loaded {len(docs)} documents from local files")
        
        # SharePoint
        if self.config['sharepoint']['enabled']:
            logger.info("Loading from SharePoint...")
            docs = self.load_sharepoint(self.config['sharepoint'])
            all_documents.extend(docs)
        
        # Google Drive
        if self.config['google_drive']['enabled']:
            logger.info("Loading from Google Drive...")
            docs = self.load_google_drive(self.config['google_drive'])
            all_documents.extend(docs)
        
        # Slack
        if self.config['slack']['enabled']:
            logger.info("Loading from Slack...")
            docs = self.load_slack(self.config['slack'])
            all_documents.extend(docs)
        
        # GitHub
        if self.config['github']['enabled']:
            logger.info("Loading from GitHub...")
            docs = self.load_github(self.config['github'])
            all_documents.extend(docs)
        
        self.documents = all_documents
        logger.info(f"Total documents loaded: {len(all_documents)}")
        
        return all_documents
    
    def save_processed_documents(self, output_path: Path = None):
        """Save processed documents to JSON file."""
        if output_path is None:
            output_path = PROCESSED_DATA_DIR / "processed_documents.jsonl"
        
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            for doc in self.documents:
                f.write(json.dumps(doc.to_dict()) + '\n')
        
        logger.info(f"Saved {len(self.documents)} documents to {output_path}")
    
    def load_processed_documents(self, input_path: Path = None) -> List[Document]:
        """Load processed documents from JSON file."""
        if input_path is None:
            input_path = PROCESSED_DATA_DIR / "processed_documents.jsonl"
        
        if not input_path.exists():
            logger.warning(f"No processed documents found at {input_path}")
            return []
        
        documents = []
        with open(input_path, 'r', encoding='utf-8') as f:
            for line in f:
                data = json.loads(line)
                doc = Document(**data)
                documents.append(doc)
        
        self.documents = documents
        logger.info(f"Loaded {len(documents)} processed documents")
        return documents
    
    def get_statistics(self) -> Dict:
        """Get statistics about loaded documents."""
        if not self.documents:
            return {}
        
        total_chars = sum(len(doc.text) for doc in self.documents)
        doc_types = {}
        sources = {}
        
        for doc in self.documents:
            doc_types[doc.doc_type] = doc_types.get(doc.doc_type, 0) + 1
            sources[doc.source] = sources.get(doc.source, 0) + 1
        
        return {
            'total_documents': len(self.documents),
            'total_characters': total_chars,
            'avg_document_length': total_chars / len(self.documents),
            'doc_types': doc_types,
            'sources': sources,
        }


# ============================================================================
# MAIN EXECUTION
# ============================================================================
if __name__ == "__main__":
    # Example usage
    print("Enterprise Data Loader")
    print("=" * 80)
    
    # Initialize loader
    loader = EnterpriseDataLoader()
    
    # Load all documents
    documents = loader.load_all_sources()
    
    # Print statistics
    stats = loader.get_statistics()
    print("\nDataset Statistics:")
    print(f"Total Documents: {stats.get('total_documents', 0)}")
    print(f"Total Characters: {stats.get('total_characters', 0):,}")
    print(f"Average Length: {stats.get('avg_document_length', 0):.2f} characters")
    print(f"\nDocument Types: {stats.get('doc_types', {})}")
    print(f"Sources: {stats.get('sources', {})}")
    
    # Save processed documents
    if documents:
        loader.save_processed_documents()
        print(f"\nProcessed documents saved to: {PROCESSED_DATA_DIR / 'processed_documents.jsonl'}")
